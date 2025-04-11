import os
import pandas as pd
import docx
import re
import spacy
import networkx as nx
from tqdm import tqdm
from concurrent.futures import ProcessPoolExecutor
import matplotlib.pyplot as plt
from typing import List, Dict, Tuple, Set, Any, Optional
import time
import logging


# Step 1: Read text from Word documents
def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text content from a Word document.

    Args:
        file_path: Path to the Word document

    Returns:
        String containing all text from the document
    """
    try:
        doc = docx.Document(file_path)
        full_text = []

        # Extract text from paragraphs
        for para in doc.paragraphs:
            full_text.append(para.text)

        # Also get text from tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)

        return "\n".join(full_text)
    except Exception as e:
        print(f"Error processing {file_path}: {e}")
        return ""


def process_documents(directory: str) -> pd.DataFrame:
    """
    Process all Word documents in a directory and return a DataFrame.

    Args:
        directory: Path to directory containing Word documents

    Returns:
        DataFrame with columns: filename, content
    """
    results = []

    # Get all docx files in the directory
    docx_files = [f for f in os.listdir(directory) if f.endswith(".docx")]

    print(f"Found {len(docx_files)} Word documents to process")

    # Process each document
    for filename in tqdm(docx_files, desc="Processing documents"):
        file_path = os.path.join(directory, filename)
        content = extract_text_from_docx(file_path)

        if content:  # Only add if we got content
            results.append({"filename": filename, "content": content})

    # Convert to DataFrame
    return pd.DataFrame(results)


# Step 2: Chunk text for encoder input
def chunk_text(text: str, chunk_size: int = 512, overlap: int = 50) -> List[str]:
    """
    Split text into overlapping chunks of specified size.

    Args:
        text: The text to chunk
        chunk_size: Maximum size of each chunk in tokens (approximated by words)
        overlap: Number of tokens to overlap between chunks

    Returns:
        List of text chunks
    """
    # Simple word-based chunking (approximation for tokens)
    words = text.split()

    if len(words) <= chunk_size:
        return [text]

    chunks = []
    start = 0

    while start < len(words):
        # Get chunk of words
        end = min(start + chunk_size, len(words))
        chunk = " ".join(words[start:end])
        chunks.append(chunk)

        # Move start position for next chunk, accounting for overlap
        start = start + chunk_size - overlap

    return chunks


def process_documents_with_chunking(
    directory: str, chunk_size: int = 512, overlap: int = 50
) -> pd.DataFrame:
    """
    Process all Word documents in a directory and chunk their content.

    Args:
        directory: Path to directory containing Word documents
        chunk_size: Maximum size of each chunk in tokens
        overlap: Number of tokens to overlap between chunks

    Returns:
        DataFrame with columns: filename, chunk_id, chunk_text
    """
    # First process all documents
    df_documents = process_documents(directory)

    # Then chunk each document
    results = []

    for _, row in tqdm(
        df_documents.iterrows(), desc="Chunking documents", total=len(df_documents)
    ):
        chunks = chunk_text(row["content"], chunk_size, overlap)

        for i, chunk in enumerate(chunks):
            results.append(
                {"filename": row["filename"], "chunk_id": i, "chunk_text": chunk}
            )

    return pd.DataFrame(results)


# Step 3: Create entity-relation mapping (knowledge graph)


def initialize_nlp():
    """
    Initialize and return the spaCy NLP pipeline.

    Returns:
        spaCy NLP pipeline with necessary components
    """
    try:
        # Load spaCy English model with the needed components
        # You may need to download this model: python -m spacy download en_core_web_lg
        nlp = spacy.load("en_core_web_lg")

        # Add the dependency parse and entity recognition components if not present
        if "parser" not in nlp.pipe_names:
            nlp.add_pipe("parser")
        if "ner" not in nlp.pipe_names:
            nlp.add_pipe("ner")

        return nlp
    except Exception as e:
        print(f"Error initializing spaCy: {e}")
        print(
            "Make sure to install spaCy and download the model: python -m spacy download en_core_web_lg"
        )
        raise


def extract_entities_and_relations(text: str, nlp) -> Tuple[List[Dict], List[Dict]]:
    """
    Extract entities and relations from text using spaCy.

    Args:
        text: Text to process
        nlp: spaCy NLP pipeline

    Returns:
        Tuple containing list of entities and list of relations
    """
    # Process the text
    doc = nlp(text)

    # Extract entities
    entities = []
    for ent in doc.ents:
        entities.append(
            {
                "id": ent.text.lower().replace(" ", "_"),
                "text": ent.text,
                "label": ent.label_,
                "start": ent.start_char,
                "end": ent.end_char,
            }
        )

    # Extract relations using dependency parsing
    relations = []

    # Focus on subject-verb-object structures
    for token in doc:
        if token.dep_ in ("nsubj", "nsubjpass") and token.head.pos_ == "VERB":
            # Found a subject-verb relation
            subject = token.text
            verb = token.head.text

            # Look for an object
            for child in token.head.children:
                if child.dep_ in ("dobj", "pobj"):
                    obj = child.text
                    relations.append(
                        {
                            "subject": subject.lower(),
                            "predicate": verb.lower(),
                            "object": obj.lower(),
                            "sentence": token.sent.text,
                        }
                    )

    return entities, relations


def build_knowledge_graph(
    df_chunks: pd.DataFrame,
) -> Tuple[nx.Graph, pd.DataFrame, pd.DataFrame]:
    """
    Build a knowledge graph from document chunks.

    Args:
        df_chunks: DataFrame with columns: filename, chunk_id, chunk_text

    Returns:
        Tuple containing:
        - NetworkX graph
        - DataFrame of entities
        - DataFrame of relations
    """
    # Initialize spaCy
    nlp = initialize_nlp()

    all_entities = []
    all_relations = []

    # Process each chunk
    for _, row in tqdm(
        df_chunks.iterrows(),
        desc="Extracting entities and relations",
        total=len(df_chunks),
    ):
        entities, relations = extract_entities_and_relations(row["chunk_text"], nlp)

        # Add source information to entities and relations
        for entity in entities:
            entity["source_file"] = row["filename"]
            entity["chunk_id"] = row["chunk_id"]

        for relation in relations:
            relation["source_file"] = row["filename"]
            relation["chunk_id"] = row["chunk_id"]

        all_entities.extend(entities)
        all_relations.extend(relations)

    # Convert to DataFrames
    df_entities = pd.DataFrame(all_entities)
    df_relations = pd.DataFrame(all_relations)

    # Remove duplicates
    if not df_entities.empty:
        df_entities.drop_duplicates(subset=["id"], inplace=True)

    if not df_relations.empty:
        df_relations.drop_duplicates(
            subset=["subject", "predicate", "object"], inplace=True
        )

    # Create NetworkX graph
    G = nx.Graph()

    # Add entities as nodes
    for _, entity in df_entities.iterrows():
        G.add_node(entity["id"], label=entity["text"], type=entity["label"])

    # Add relations as edges
    for _, relation in df_relations.iterrows():
        subject_id = relation["subject"].replace(" ", "_")
        object_id = relation["object"].replace(" ", "_")

        # Only add edge if both nodes exist
        if subject_id in G.nodes and object_id in G.nodes:
            G.add_edge(subject_id, object_id, relation=relation["predicate"])

    return G, df_entities, df_relations


def visualize_graph(G: nx.Graph, max_nodes: int = 50):
    """
    Visualize the knowledge graph.

    Args:
        G: NetworkX graph
        max_nodes: Maximum number of nodes to display
    """
    # Take a subset of the graph if it's too large
    if len(G.nodes) > max_nodes:
        nodes = list(G.nodes)[:max_nodes]
        G = G.subgraph(nodes)

    plt.figure(figsize=(12, 12))
    pos = nx.spring_layout(G)

    nx.draw(
        G,
        pos,
        with_labels=True,
        node_color="skyblue",
        node_size=1500,
        edge_color="gray",
        linewidths=1,
        font_size=10,
        font_weight="bold",
    )

    plt.title(f"Knowledge Graph (showing {len(G.nodes)} nodes)")
    plt.tight_layout()
    plt.savefig("knowledge_graph.png", format="PNG", dpi=300)
    plt.show()


def extract_game_elements(
    G: nx.Graph, df_entities: pd.DataFrame, df_relations: pd.DataFrame
) -> Dict:
    """
    Extract specific game elements from the knowledge graph.

    Args:
        G: NetworkX graph
        df_entities: DataFrame of entities
        df_relations: DataFrame of relations

    Returns:
        Dictionary containing game elements
    """
    game_elements = {
        "locations": set(),
        "characters": set(),
        "items": set(),
        "actions": set(),
        "character_relations": [],
        "location_connections": [],
    }

    # Extract locations (GPE, LOC, FAC entities)
    location_types = ["GPE", "LOC", "FAC"]
    for _, entity in df_entities.iterrows():
        if entity["label"] in location_types:
            game_elements["locations"].add(entity["text"])

    # Extract characters (PERSON entities)
    for _, entity in df_entities.iterrows():
        if entity["label"] == "PERSON":
            game_elements["characters"].add(entity["text"])

    # Extract items (potential objects)
    for _, relation in df_relations.iterrows():
        if relation["predicate"] in ["have", "carry", "hold", "own", "possess"]:
            game_elements["items"].add(relation["object"])

    # Extract actions (verbs in relations)
    for _, relation in df_relations.iterrows():
        game_elements["actions"].add(relation["predicate"])

    # Extract character relations
    for _, relation in df_relations.iterrows():
        if (
            relation["subject"] in game_elements["characters"]
            and relation["object"] in game_elements["characters"]
        ):
            game_elements["character_relations"].append(relation)

    # Convert sets to lists for JSON serialization
    game_elements["locations"] = list(game_elements["locations"])
    game_elements["characters"] = list(game_elements["characters"])
    game_elements["items"] = list(game_elements["items"])
    game_elements["actions"] = list(game_elements["actions"])

    return game_elements


# Example usage
# nlp = initialize_nlp()
# df_chunks = process_documents_with_chunking("/path/to/documents")
# G, df_entities, df_relations = build_knowledge_graph(df_chunks)
# print(f"Built knowledge graph with {len(G.nodes)} nodes and {len(G.edges)} edges")
# visualize_graph(G)
# game_elements = extract_game_elements(G, df_entities, df_relations)
# print(f"Extracted {len(game_elements['locations'])} locations, {len(game_elements['characters'])} characters")


def main(
    documents_dir: str,
    output_dir: str = "output",
    chunk_size: int = 512,
    overlap: int = 50,
    output_name: str = None,
    process_quests: bool = True,
    use_llm: bool = True,
):
    """
    Run the complete pipeline from document processing to game element extraction.

    Args:
        documents_dir: Directory containing Word documents
        output_dir: Directory to save output files
        chunk_size: Maximum size of each chunk in tokens
        overlap: Number of tokens to overlap between chunks
        output_name: Optional name for the output subfolder (defaults to input directory name if not provided)
    """
    # Create a named output directory based on input directory if not specified
    if not output_name:
        # Extract the last directory name from the documents_dir path
        output_name = os.path.basename(os.path.normpath(documents_dir))
        # If the extracted name is 'documents' (the default root), use a timestamp instead
        if output_name == "documents":
            timestamp = time.strftime("%Y%m%d_%H%M%S")
            output_name = f"world_{timestamp}"

    world_output_dir = os.path.join(output_dir, output_name)
    os.makedirs(world_output_dir, exist_ok=True)

    print("Step 1: Processing documents...")
    df_chunks = process_documents_with_chunking(documents_dir, chunk_size, overlap)
    df_chunks.to_csv(os.path.join(world_output_dir, "document_chunks.csv"), index=False)
    print(
        f"Created {len(df_chunks)} chunks from {df_chunks['filename'].nunique()} documents"
    )

    print("\nStep 2: Building knowledge graph...")
    G, df_entities, df_relations = build_knowledge_graph(df_chunks)
    print(f"Built knowledge graph with {len(G.nodes)} nodes and {len(G.edges)} edges")

    # Save results
    df_entities.to_csv(os.path.join(world_output_dir, "entities.csv"), index=False)
    df_relations.to_csv(os.path.join(world_output_dir, "relations.csv"), index=False)

    # Save graph
    nx.write_gexf(G, os.path.join(world_output_dir, "knowledge_graph.gexf"))

    print("\nStep 3: Visualizing graph...")
    plt.figure(figsize=(12, 12))
    visualize_graph(G)
    plt.savefig(os.path.join(world_output_dir, "graph_sample.png"))
    plt.close()

    print("\nStep 4: Extracting game elements...")
    game_elements = extract_game_elements(G, df_entities, df_relations)

    # Save game elements as CSV files
    pd.DataFrame(game_elements["locations"], columns=["location"]).to_csv(
        os.path.join(world_output_dir, "game_locations.csv"), index=False
    )

    pd.DataFrame(game_elements["characters"], columns=["character"]).to_csv(
        os.path.join(world_output_dir, "game_characters.csv"), index=False
    )

    pd.DataFrame(game_elements["items"], columns=["item"]).to_csv(
        os.path.join(world_output_dir, "game_items.csv"), index=False
    )

    pd.DataFrame(game_elements["actions"], columns=["action"]).to_csv(
        os.path.join(world_output_dir, "game_actions.csv"), index=False
    )

    pd.DataFrame(game_elements["character_relations"]).to_csv(
        os.path.join(world_output_dir, "game_character_relations.csv"), index=False
    )

    print(f"\nExtracted game elements:")
    print(f"- {len(game_elements['locations'])} locations")
    print(f"- {len(game_elements['characters'])} characters")
    print(f"- {len(game_elements['items'])} items")
    print(f"- {len(game_elements['actions'])} actions")
    print(f"- {len(game_elements['character_relations'])} character relationships")

    # Process quest documents if enabled
    if process_quests:
        # Extract world name from the output directory
        current_world_name = os.path.basename(world_output_dir)
        process_quest_documents(documents_dir, output_dir, current_world_name, use_llm)

    print(f"\nAll results saved to {world_output_dir}/")

    return world_output_dir


def list_document_folders(base_dir="data/documents"):
    """List available document folders for processing."""
    # Create the directory if it doesn't exist
    os.makedirs(base_dir, exist_ok=True)

    print("\nAvailable document folders:")
    print("-" * 80)

    # Get all subdirectories
    folders = []
    for item in os.listdir(base_dir):
        item_path = os.path.join(base_dir, item)
        if os.path.isdir(item_path):
            docx_count = len([f for f in os.listdir(item_path) if f.endswith(".docx")])
            folders.append((item, os.path.join(base_dir, item), docx_count))

    # Add the root documents directory
    docx_count = len([f for f in os.listdir(base_dir) if f.endswith(".docx")])
    if docx_count > 0:
        folders.insert(0, ("root", base_dir, docx_count))

    if not folders:
        print("No document folders found with .docx files.")
        print(f"Create a subfolder in {base_dir} and add .docx files to it.")
        return

    # Print folders
    for name, path, count in folders:
        print(f"Name: {name}")
        print(f"Path: {path}")
        print(f"Document count: {count} .docx files")
        print("-" * 80)


def process_quest_documents(
    documents_dir: str,
    output_dir: str,
    world_name: str,
    use_llm: bool = True,
    debug_mode: bool = True,
) -> None:
    """
    Process quest documents in a directory.

    Args:
        documents_dir: Directory containing documents
        output_dir: Directory to save output files
        world_name: Name of the world
        use_llm: Whether to use an LLM for quest parsing
        debug_mode: Whether to enable debug mode for quest parsing
    """
    try:
        # Import here to avoid circular imports
        import sys
        import os
        import logging
        import json

        logger = logging.getLogger(__name__)
        logger.info(f"Current working directory: {os.getcwd()}")
        logger.info(f"Python path: {sys.path}")

        try:
            # Try importing directly (when script is run directly)
            logger.info("Attempting direct import of DocumentQuestProcessor")
            from document_quest_processor import DocumentQuestProcessor

            logger.info("Direct import successful")
        except ImportError as e:
            logger.warning(f"Direct import failed: {e}")
            # If that fails, try importing as a module
            try:
                logger.info("Attempting module import of DocumentQuestProcessor")
                from src.document_quest_processor import DocumentQuestProcessor

                logger.info("Module import successful")
            except ImportError as e:
                logger.warning(f"Module import failed: {e}")
                # If both fail, try a relative import
                try:
                    logger.info("Attempting relative import of DocumentQuestProcessor")
                    from .document_quest_processor import DocumentQuestProcessor

                    logger.info("Relative import successful")
                except ImportError as e:
                    logger.error(f"All import attempts failed: {e}")
                    # Check if the file exists
                    quest_processor_path = os.path.join(
                        os.path.dirname(__file__), "document_quest_processor.py"
                    )
                    if os.path.exists(quest_processor_path):
                        logger.info(f"File exists at {quest_processor_path}")
                    else:
                        logger.error(f"File does not exist at {quest_processor_path}")
                    # If all fail, raise the error to be caught by the outer try/except
                    raise ImportError("Could not import DocumentQuestProcessor")

        # Find quest documents (files ending with quest.docx, case insensitive)
        quest_docs = []
        for filename in os.listdir(documents_dir):
            if filename.lower().endswith("quest.docx"):
                quest_docs.append(os.path.join(documents_dir, filename))

        if not quest_docs:
            print("No quest documents found.")
            return

        print(f"\nFound {len(quest_docs)} quest document(s). Processing...")
        for doc in quest_docs:
            print(f"  - {os.path.basename(doc)}")

        # Select LLM client if requested
        llm_client = None
        if use_llm:
            try:
                from llm_clients import select_llm_client

                llm_client = select_llm_client(interactive=True, debug_mode=debug_mode)
                print(f"Using {llm_client.name} for quest parsing")
            except ImportError:
                print(
                    "LLM client module not available. Using rule-based quest parsing."
                )
                from llm_clients import RuleBasedClient

                llm_client = RuleBasedClient(debug_mode=debug_mode)
            except Exception as e:
                print(
                    f"Error selecting LLM client: {e}. Using rule-based quest parsing."
                )
                from llm_clients import RuleBasedClient

                llm_client = RuleBasedClient(debug_mode=debug_mode)
        else:
            # Use rule-based client with debug mode
            from llm_clients import RuleBasedClient

            llm_client = RuleBasedClient(debug_mode=debug_mode)
            print("Using rule-based quest parsing (no LLM)")

        # Create a temporary directory for quest documents
        import tempfile

        with tempfile.TemporaryDirectory() as temp_dir:
            # Copy quest documents to temp directory
            import shutil

            for doc in quest_docs:
                shutil.copy2(doc, temp_dir)

            # Process quest documents
            processor = DocumentQuestProcessor(llm_client)
            quests, df_quests = processor.process_documents_for_quests(
                temp_dir, output_dir, world_name
            )

            # Output quests as JSON for debugging
            if debug_mode and quests:
                debug_output_dir = os.path.join(output_dir, "debug")
                os.makedirs(debug_output_dir, exist_ok=True)
                debug_file = os.path.join(
                    debug_output_dir, f"{world_name}_quests_debug.json"
                )

                # Convert Quest objects to dictionaries for JSON serialization
                quest_dicts = []
                for quest in quests:
                    try:
                        # If it's already a dict, use it directly
                        if isinstance(quest, dict):
                            quest_dicts.append(quest)
                        # Otherwise, convert Quest object to dict
                        else:
                            quest_dict = {
                                "quest_id": quest.quest_id,
                                "title": quest.title,
                                "description": quest.description,
                                "quest_type": str(quest.quest_type),
                                "giver": quest.giver
                                if hasattr(quest, "giver")
                                else None,
                                "reward": quest.reward
                                if hasattr(quest, "reward")
                                else None,
                            }

                            # Add type-specific attributes
                            if hasattr(quest, "target_location"):
                                quest_dict["target_location"] = quest.target_location
                            if hasattr(quest, "item_to_fetch"):
                                quest_dict["item_to_fetch"] = quest.item_to_fetch
                            if hasattr(quest, "recipient"):
                                quest_dict["recipient"] = quest.recipient
                            if hasattr(quest, "target_enemy"):
                                quest_dict["target_enemy"] = quest.target_enemy
                            if hasattr(quest, "information"):
                                quest_dict["information"] = quest.information
                            if hasattr(quest, "source"):
                                quest_dict["source"] = quest.source
                            if hasattr(quest, "target"):
                                quest_dict["target"] = quest.target
                            if hasattr(quest, "difficulty"):
                                quest_dict["difficulty"] = quest.difficulty
                            if hasattr(quest, "time_limit"):
                                quest_dict["time_limit"] = quest.time_limit

                            # Handle success and failure consequences
                            if (
                                hasattr(quest, "success_consequences")
                                and quest.success_consequences
                            ):
                                success_consequences = []
                                for consequence in quest.success_consequences:
                                    success_consequences.append(
                                        {
                                            "consequence_type": consequence.consequence_type,
                                            "target": consequence.target,
                                            "parameters": consequence.parameters,
                                        }
                                    )
                                quest_dict["success_consequences"] = (
                                    success_consequences
                                )

                            if (
                                hasattr(quest, "failure_consequences")
                                and quest.failure_consequences
                            ):
                                failure_consequences = []
                                for consequence in quest.failure_consequences:
                                    failure_consequences.append(
                                        {
                                            "consequence_type": consequence.consequence_type,
                                            "target": consequence.target,
                                            "parameters": consequence.parameters,
                                        }
                                    )
                                quest_dict["failure_consequences"] = (
                                    failure_consequences
                                )

                            # Add notes if available
                            if hasattr(quest, "notes") and quest.notes:
                                quest_dict["notes"] = quest.notes

                            quest_dicts.append(quest_dict)
                    except Exception as e:
                        logger.error(f"Error converting quest to dict: {e}")

                # Save the quest dictionaries as JSON
                with open(debug_file, "w") as f:
                    json.dump(quest_dicts, f, indent=2)
                print(f"Saved debug quest data to {debug_file}")

            print(f"Extracted {len(quests)} quests from documents")
            if len(quests) > 0:
                for quest in quests:
                    if isinstance(quest, dict):
                        title = quest.get("title", "Unnamed quest")
                        quest_type = quest.get("quest_type", "Unknown type")
                    else:
                        title = (
                            quest.title if hasattr(quest, "title") else "Unnamed quest"
                        )
                        quest_type = (
                            str(quest.quest_type)
                            if hasattr(quest, "quest_type")
                            else "Unknown type"
                        )
                    print(f"  - {title} ({quest_type})")
            else:
                print("No quests were extracted. Check the quest document format.")

    except ImportError:
        print("Quest processing module not available. Skipping quest processing.")
    except Exception as e:
        print(f"Error processing quest documents: {e}")


def print_help_message():
    """
    Print a detailed help message explaining how to use the document processor.
    """
    help_text = """
=== GraphRAG Document Processor Help ===

The document processor is used to extract game elements from Word documents and create a world for the GraphRAG text adventure game.

Available Commands:
  list                    List available document folders in data/documents/
  process                 Process documents into a world
  help                    Show this help message

Process Command Usage:
  python document_processor.py process --folder <world_name>

Example:
  python document_processor.py process --folder fantasy_world

This will:
  1. Process all .docx files in data/documents/fantasy_world/
  2. Extract locations, characters, items, and relationships
  3. Automatically process any files ending with 'quest.docx' to extract quests
  4. Save all data to data/output/fantasy_world/

Options:
  --folder <name>         Name of the document folder to process (subfolder of data/documents)
  --documents_dir <path>  Alternative: Full path to directory containing Word documents
  --output_dir <path>     Directory to save output files (default: data/output)
  --output_name <name>    Custom name for the output world directory
  --chunk_size <int>      Maximum chunk size in tokens (default: 512)
  --overlap <int>         Overlap between chunks in tokens (default: 50)
  --skip_quests           Skip processing quest documents
  --no-llm                Disable LLM for quest parsing (use rule-based parsing)

Quest Documents:
  Any document whose filename ends with 'quest.docx' (case insensitive) will be
  automatically processed to extract quest information. Quests will be integrated
  into the world knowledge graph and saved to quests.json in the output directory.

  By default, the processor will prompt you to select an LLM for quest parsing.
  Available options may include:
  - Rule-Based (No LLM): Simple pattern matching, always available
  - Local LLM: If you have a local model installed in the models/llm directory
  - OpenAI GPT-3.5: If you have an OpenAI API key set

Examples:
  # List available document folders
  python document_processor.py list

  # Process a specific folder
  python document_processor.py process --folder fantasy_world

  # Process with custom output name
  python document_processor.py process --folder fantasy_world --output_name my_fantasy_world

  # Skip quest processing
  python document_processor.py process --folder fantasy_world --skip_quests

  # Use rule-based quest parsing (no LLM)
  python document_processor.py process --folder fantasy_world --no-llm

  # Use a custom documents directory
  python document_processor.py process --documents_dir /path/to/my/documents
"""
    print(help_text)


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(
        description="Process Word documents for a GraphRAG-based text adventure game"
    )

    # Create subparsers for different commands
    subparsers = parser.add_subparsers(dest="command", help="Command to execute")

    # Create the 'list' command
    list_parser = subparsers.add_parser("list", help="List available document folders")

    # Create the 'help' command
    help_parser = subparsers.add_parser("help", help="Show detailed help information")

    # Create the 'process' command
    process_parser = subparsers.add_parser(
        "process", help="Process documents into a world"
    )
    process_parser.add_argument(
        "--documents_dir", help="Directory containing Word documents"
    )
    process_parser.add_argument(
        "--folder",
        help="Name of the document folder to process (subfolder of data/documents)",
    )
    process_parser.add_argument(
        "--output_dir", default="data/output", help="Directory to save output files"
    )
    process_parser.add_argument(
        "--chunk_size", type=int, default=512, help="Maximum chunk size in tokens"
    )
    process_parser.add_argument(
        "--overlap", type=int, default=50, help="Overlap between chunks in tokens"
    )
    process_parser.add_argument(
        "--output_name", help="Optional name for the output world directory"
    )
    process_parser.add_argument(
        "--skip_quests", action="store_true", help="Skip processing quest documents"
    )
    process_parser.add_argument(
        "--no-llm",
        action="store_true",
        help="Disable LLM for quest parsing (use rule-based parsing)",
    )

    args = parser.parse_args()

    # Handle commands
    if args.command == "list":
        list_document_folders()
    elif args.command == "help":
        print_help_message()
    elif args.command == "process":
        # Determine the documents directory
        documents_dir = args.documents_dir
        if args.folder and not documents_dir:
            documents_dir = os.path.join("data/documents", args.folder)

        if not documents_dir:
            print("Error: You must specify either --documents_dir or --folder")
            import sys

            sys.exit(1)

        # Process the documents
        main(
            documents_dir,
            args.output_dir,
            args.chunk_size,
            args.overlap,
            args.output_name,
            not args.skip_quests,  # Process quests unless skip_quests is True
            not getattr(args, "no_llm", False),  # Use LLM unless no-llm is True
        )
    else:
        # If no command is provided, show help
        print(
            "\nUse 'python document_processor.py help' for detailed usage instructions."
        )
        parser.print_help()
